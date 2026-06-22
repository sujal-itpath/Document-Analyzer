import os
import pdfplumber
import docx
import logging
from unstructured.partition.auto import partition
from langchain_core.documents import Document
from typing import List

logger = logging.getLogger(__name__)

class DocumentProcessor:
    @staticmethod
    def process_document(file_path: str) -> List[Document]:
        """ 
        Process a document using OCR and table extraction.
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return DocumentProcessor._process_pdf(file_path)
        elif ext in ['.png', '.jpg', '.jpeg']:
            return DocumentProcessor._process_image(file_path)
        elif ext == '.docx':
            return DocumentProcessor._process_docx(file_path)
        else:
            # Fallback to unstructured for other types
            elements = partition(filename=file_path)
            docs = []
            for el in elements:
                docs.append(Document(
                    page_content=str(el),
                    metadata={"source": file_path, "type": str(type(el))}
                ))
            return docs

    @staticmethod
    def _process_pdf(file_path: str) -> List[Document]:
        documents = []
        extracted_text_length = 0
        
        # 1. Extract text and tables using pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                # Extract text
                text = page.extract_text()
                if text:
                    extracted_text_length += len(text.strip())
                    # Include page number in content for better RAG context
                    documents.append(Document(
                        page_content=f"[Page {i + 1}]\n{text}",
                        metadata={"source": file_path, "page": i + 1, "content_type": "text"}
                    ))
                
                # Extract tables
                tables = page.extract_tables()
                for table_idx, table in enumerate(tables):
                    if table:
                        # Convert table to markdown or CSV-like string
                        table_str = DocumentProcessor._format_table(table)
                        documents.append(Document(
                            page_content=f"Table on Page {i + 1}:\n{table_str}",
                            metadata={"source": file_path, "page": i + 1, "content_type": "table", "table_index": table_idx}
                        ))
        
        # 2. If no text was found, or it's very short (likely a scanned PDF), try OCR.
        # Increased threshold to 800 characters to be safer.
        if not documents or extracted_text_length < 800:
            logger.info(
                "Low text content (%d chars). Attempting OCR for %s.",
                extracted_text_length,
                file_path,
            )
            elements = partition(filename=file_path, strategy="ocr_only")
            for el in elements:
                documents.append(Document(
                    page_content=str(el),
                    metadata={"source": file_path, "content_type": "ocr"}
                ))
                
        return documents

    @staticmethod
    def _process_image(file_path: str) -> List[Document]:
        # Use unstructured for images (OCR)
        elements = partition(filename=file_path, strategy="hi_res")
        docs = []
        for el in elements:
            docs.append(Document(
                page_content=str(el),
                metadata={"source": file_path, "content_type": "ocr"}
            ))
        return docs

    @staticmethod
    def _process_docx(file_path: str) -> List[Document]:
        doc = docx.Document(file_path)
        documents = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                documents.append(Document(
                    page_content=text,
                    metadata={"source": file_path, "content_type": "text"}
                ))
        return documents

    @staticmethod
    def _format_table(table: List[List[str]]) -> str:
        """Formats a list of lists into a markdown-like table string."""
        if not table:
            return ""
        
        # Clean none values
        cleaned_table = [[str(cell) if cell is not None else "" for cell in row] for row in table]
        
        # Simple string representation
        rows = [" | ".join(row) for row in cleaned_table]
        return "\n".join(rows)
